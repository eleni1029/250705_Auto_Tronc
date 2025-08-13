#!/usr/bin/env python3
"""
Word文件章節拆分器
掃描Word文件中的章節標題，並根據章節將文件拆分成多個檔案
"""

import os
import re
import docx
import logging
from pathlib import Path
from docx.document import Document
from collections import Counter

class WordSplitter:
    def __init__(self, logger=None):
        self.logger = logger or self._setup_default_logger()
        
        # 章節關鍵字模式
        self.chapter_patterns = [
            r'第(.{1,3})[章節篇講堂課單元部分週周]',  # 第一章、第二節、第三週等
            r'第(\d{1,2})[章節篇講堂課單元部分週周]',  # 第1章、第12節、第3週等
            r'(\d{1,2})\s*[\.、]\s*[章節篇講堂課單元部分週周]',  # 1. 章節、2、章節等
        ]
        
        # 中文數字轉換 (1-99)
        self.chinese_numbers = self._build_chinese_numbers()
        
    def _build_chinese_numbers(self):
        """建立1-99的中文數字對照表"""
        chinese_nums = {}
        
        # 1-9
        basic_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九']
        for i, num in enumerate(basic_nums, 1):
            chinese_nums[num] = str(i)
        
        # 10
        chinese_nums['十'] = '10'
        
        # 11-19
        for i in range(1, 10):
            chinese_nums[f'十{basic_nums[i-1]}'] = str(10 + i)
        
        # 20-99
        tens = ['二', '三', '四', '五', '六', '七', '八', '九']
        for i, ten in enumerate(tens, 2):
            # X十 (20, 30, 40, ..., 90)
            chinese_nums[f'{ten}十'] = str(i * 10)
            
            # X十Y (21-29, 31-39, ..., 91-99)
            for j in range(1, 10):
                chinese_nums[f'{ten}十{basic_nums[j-1]}'] = str(i * 10 + j)
        
        return chinese_nums
    
    def _setup_default_logger(self):
        """設置預設logger"""
        logger = logging.getLogger('WordSplitter')
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
            logger.setLevel(logging.INFO)
        return logger
    
    def clean_text(self, text):
        """清除全形與半形空格"""
        if not text:
            return ""
        # 清除全形空格（\u3000）和半形空格
        return re.sub(r'[\s\u3000]+', '', text.strip())
    
    def is_chapter_title(self, text):
        """判斷是否為章節標題"""
        if not text or len(text) > 20:
            return False
            
        cleaned_text = self.clean_text(text)
        if len(cleaned_text) > 20:
            return False
            
        # 檢查是否符合章節模式
        for pattern in self.chapter_patterns:
            if re.search(pattern, cleaned_text):
                return True
                
        return False
    
    def extract_chapter_info(self, text):
        """提取章節資訊"""
        cleaned_text = self.clean_text(text)
        
        for pattern in self.chapter_patterns:
            match = re.search(pattern, cleaned_text)
            if match:
                chapter_num = match.group(1)
                # 如果是中文數字，轉換為阿拉伯數字
                if chapter_num in self.chinese_numbers:
                    chapter_num = self.chinese_numbers[chapter_num]
                
                return {
                    'number': chapter_num,
                    'title': cleaned_text,
                    'original_text': text.strip()
                }
        
        return None
    
    def scan_document_structure(self, doc_path):
        """掃描文檔結構，找出章節分割點"""
        try:
            doc = docx.Document(doc_path)
            chapters = []
            
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                    
                if self.is_chapter_title(text):
                    chapter_info = self.extract_chapter_info(text)
                    if chapter_info:
                        chapter_info['paragraph_index'] = i
                        chapters.append(chapter_info)
                        self.logger.info(f"找到章節: {chapter_info['title']} (段落 {i})")
            
            return chapters
            
        except Exception as e:
            self.logger.error(f"掃描文檔結構失敗 {doc_path}: {e}")
            return []
    
    def split_document_by_chapters(self, doc_path, output_dir, base_name):
        """根據章節拆分文檔"""
        try:
            doc = docx.Document(doc_path)
            chapters = self.scan_document_structure(doc_path)
            
            if not chapters:
                self.logger.info(f"未找到章節標題，不進行拆分: {doc_path}")
                return False
            
            self.logger.info(f"找到 {len(chapters)} 個章節，開始拆分")
            
            # 確保輸出目錄存在
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # 記錄已使用的檔名，避免重複
            used_names = set()
            created_files = []
            
            for i, chapter in enumerate(chapters):
                # 創建新文檔
                new_doc = docx.Document()
                
                # 確定段落範圍
                start_idx = chapter['paragraph_index']
                end_idx = chapters[i + 1]['paragraph_index'] if i + 1 < len(chapters) else len(doc.paragraphs)
                
                # 複製段落內容
                content_lines = 0
                paragraphs_with_content = []
                
                for para_idx in range(start_idx, end_idx):
                    if para_idx < len(doc.paragraphs):
                        para = doc.paragraphs[para_idx]
                        if para.text.strip():  # 只記錄有內容的段落
                            paragraphs_with_content.append(para)
                            content_lines += 1
                
                # 如果只有一行內容，跳過不生成檔案
                if content_lines <= 1:
                    self.logger.info(f"跳過章節 '{chapter['title']}' （只有 {content_lines} 行內容，可能是目錄）")
                    continue
                
                # 新增段落到文檔
                for para in paragraphs_with_content:
                    new_para = new_doc.add_paragraph(para.text)
                    # 嘗試保持原有格式
                    try:
                        new_para.style = para.style
                    except:
                        pass
                
                # 生成檔案 (已經檢查過內容數量)
                # 生成檔名
                clean_title = re.sub(r'[<>:"/\\|?*]', '_', chapter['title'])
                file_name = f"{base_name}_{clean_title}"
                
                # 處理重複檔名
                original_file_name = file_name
                counter = 1
                while file_name in used_names:
                    file_name = f"{original_file_name}_{counter}"
                    counter += 1
                
                used_names.add(file_name)
                output_path = output_dir / f"{file_name}.docx"
                
                # 保存檔案
                new_doc.save(output_path)
                created_files.append(output_path)
                self.logger.info(f"創建章節檔案: {output_path}")
            
            self.logger.info(f"拆分完成，共創建 {len(created_files)} 個檔案")
            return True
            
        except Exception as e:
            self.logger.error(f"拆分文檔失敗 {doc_path}: {e}")
            return False
    
    def analyze_document(self, doc_path):
        """分析文檔是否需要拆分"""
        chapters = self.scan_document_structure(doc_path)
        
        analysis = {
            'has_chapters': len(chapters) > 0,
            'chapter_count': len(chapters),
            'chapters': chapters,
            'needs_splitting': len(chapters) > 1
        }
        
        return analysis

def main():
    """測試功能"""
    import sys
    
    if len(sys.argv) != 2:
        print("用法: python sub_word_splitter.py <word文件路徑>")
        return
    
    doc_path = sys.argv[1]
    if not os.path.exists(doc_path):
        print(f"檔案不存在: {doc_path}")
        return
    
    splitter = WordSplitter()
    
    # 分析文檔
    analysis = splitter.analyze_document(doc_path)
    print(f"章節分析結果:")
    print(f"  包含章節: {analysis['has_chapters']}")
    print(f"  章節數量: {analysis['chapter_count']}")
    print(f"  需要拆分: {analysis['needs_splitting']}")
    
    if analysis['has_chapters']:
        print("  找到的章節:")
        for chapter in analysis['chapters']:
            print(f"    - {chapter['title']}")
    
    # 如果需要拆分，執行拆分
    if analysis['needs_splitting']:
        output_dir = Path(doc_path).parent / "split_chapters"
        base_name = Path(doc_path).stem
        
        print(f"\n開始拆分到: {output_dir}")
        success = splitter.split_document_by_chapters(doc_path, output_dir, base_name)
        
        if success:
            print("拆分完成!")
        else:
            print("拆分失敗!")

if __name__ == "__main__":
    main()